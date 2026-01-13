#!/usr/bin/env python3
"""
PDF Editor Pro v4.0 - Professional PDF Editing Suite
A comprehensive Adobe Acrobat Pro alternative with modern UI

Features:
- Professional dark theme UI
- Ribbon-style toolbar with tool groups
- Multi-tab document interface
- Search within documents
- Bookmarks/Outline navigation
- Comments & Sticky Notes
- Stamps library
- Watermarks, Headers & Footers
- Bates numbering
- Export to Word/Images
- OCR with invisible text layer
- Form filling
- Merge, Split, Compress
- Password protection

Auto-installs all dependencies on first run.
"""

import sys
import subprocess
import os
import platform
import urllib.request
import shutil
import tempfile
import json
from pathlib import Path
from datetime import datetime

# ============================================================================
# AUTO-INSTALLER
# ============================================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TESSERACT_DIR = os.path.join(SCRIPT_DIR, "tesseract_ocr")
CONFIG_DIR = os.path.join(SCRIPT_DIR, "pdf_editor_config")
TESSERACT_VERSION = "5.5.0"
TESSERACT_DATE = "20241111"
TESSERACT_URL = f"https://github.com/tesseract-ocr/tesseract/releases/download/{TESSERACT_VERSION}/tesseract-ocr-w64-setup-{TESSERACT_VERSION}.{TESSERACT_DATE}.exe"

def get_tesseract_path():
    if platform.system() == "Windows":
        for path in [
            os.path.join(TESSERACT_DIR, "tesseract.exe"),
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.exists(path):
                return path
    return None

def download_file(url, dest_path, desc="Downloading"):
    print(f"  {desc}...")
    try:
        opener = urllib.request.build_opener()
        opener.addheaders = [('User-Agent', 'Mozilla/5.0')]
        urllib.request.install_opener(opener)
        def hook(b, bs, ts):
            if ts > 0:
                pct = min(100, b * bs * 100 // ts)
                print(f"\r    [{'█' * (pct//3)}{'░' * (33-pct//3)}] {pct}%", end="", flush=True)
        urllib.request.urlretrieve(url, dest_path, hook)
        print()
        return True
    except Exception as e:
        print(f"\n    Error: {e}")
        return False

def install_tesseract_windows():
    print("\n  Installing Tesseract OCR...")
    os.makedirs(TESSERACT_DIR, exist_ok=True)
    temp_dir = tempfile.mkdtemp()
    installer = os.path.join(temp_dir, "setup.exe")
    try:
        if download_file(TESSERACT_URL, installer, "Downloading Tesseract OCR (~70MB)"):
            print("    Running installer...")
            subprocess.run([installer, "/S", f"/D={TESSERACT_DIR}"], capture_output=True, timeout=300)
            exe = os.path.join(TESSERACT_DIR, "tesseract.exe")
            if os.path.exists(exe):
                print("    ✓ Tesseract installed")
                shutil.rmtree(temp_dir, ignore_errors=True)
                return exe
    except:
        pass
    shutil.rmtree(temp_dir, ignore_errors=True)
    return None

def pip_install(pkg):
    for method in [
        [sys.executable, "-m", "pip", "install", pkg, "-q"],
        [sys.executable, "-m", "pip", "install", pkg, "--break-system-packages", "-q"],
        [sys.executable, "-m", "pip", "install", pkg, "--user", "-q"],
    ]:
        try:
            if subprocess.run(method, capture_output=True, timeout=120).returncode == 0:
                return True
        except:
            pass
    return False

def _try_import(name):
    try:
        __import__(name)
        return True
    except:
        return False

def check_and_install_dependencies():
    required = {'PIL': 'Pillow', 'fitz': 'PyMuPDF'}
    optional = {'pytesseract': 'pytesseract', 'docx': 'python-docx'}
    missing_req = [p for i, p in required.items() if not _try_import(i)]
    missing_opt = [p for i, p in optional.items() if not _try_import(i)]
    tesseract_needed = get_tesseract_path() is None
    
    if missing_req or missing_opt or tesseract_needed:
        print("\n╔══════════════════════════════════════════════════════════╗")
        print("║         PDF Editor Pro v4.0 - First Run Setup            ║")
        print("╚══════════════════════════════════════════════════════════╝\n")
        for pkg in missing_req + missing_opt:
            print(f"  Installing {pkg}...", end=" ", flush=True)
            print("✓" if pip_install(pkg) else "⚠")
        if tesseract_needed and platform.system() == "Windows":
            install_tesseract_windows()
        print("\n  Setup complete! Starting PDF Editor Pro...\n")
    
    for i, p in required.items():
        if not _try_import(i):
            print(f"ERROR: {p} required. Run: pip install {p}")
            sys.exit(1)
    
    if (path := get_tesseract_path()):
        os.environ["TESSERACT_CMD"] = path

check_and_install_dependencies()

# ============================================================================
# IMPORTS
# ============================================================================

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
from PIL import Image, ImageTk, ImageDraw
import fitz
import io
import threading
import math
from dataclasses import dataclass
from typing import Optional, List, Tuple, Dict, Callable, Any
from enum import Enum, auto
from collections import deque

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ============================================================================
# THEME - Professional Dark UI
# ============================================================================

class Theme:
    # Backgrounds - Layered depth
    BG_DARK = "#0d0d0d"
    BG_PRIMARY = "#161616"
    BG_SECONDARY = "#1e1e1e"
    BG_TERTIARY = "#262626"
    BG_ELEVATED = "#2d2d2d"
    BG_HOVER = "#363636"
    BG_ACTIVE = "#404040"
    BG_INPUT = "#1a1a1a"
    BG_CANVAS = "#525252"
    
    # Foregrounds
    FG_PRIMARY = "#f5f5f5"
    FG_SECONDARY = "#a0a0a0"
    FG_MUTED = "#6b6b6b"
    FG_DISABLED = "#4a4a4a"
    
    # Accents
    ACCENT = "#2563eb"
    ACCENT_LIGHT = "#3b82f6"
    ACCENT_DARK = "#1d4ed8"
    ACCENT_MUTED = "#1e3a5f"
    
    # Status colors
    SUCCESS = "#10b981"
    WARNING = "#f59e0b"
    DANGER = "#ef4444"
    INFO = "#06b6d4"
    
    # Borders
    BORDER_DARK = "#1a1a1a"
    BORDER_LIGHT = "#333333"
    BORDER_FOCUS = "#2563eb"
    
    # Special
    SELECTION = "#2563eb"
    HIGHLIGHT = "#fbbf24"
    SHADOW = "#000000"
    
    # Typography
    FONT_FAMILY = "Segoe UI"
    FONT_MONO = "Consolas"
    FONT_SIZE_XS = 9
    FONT_SIZE_SM = 10
    FONT_SIZE_MD = 11
    FONT_SIZE_LG = 12
    FONT_SIZE_XL = 14
    FONT_SIZE_XXL = 18
    
    # Spacing
    PAD_XS = 2
    PAD_SM = 4
    PAD_MD = 8
    PAD_LG = 12
    PAD_XL = 16
    PAD_XXL = 24
    
    # Sizing
    TOOLBAR_HEIGHT = 90
    SIDEBAR_WIDTH = 200
    STATUSBAR_HEIGHT = 28
    TAB_HEIGHT = 36
    BUTTON_HEIGHT = 32
    ICON_SIZE = 20

# Predefined stamps
BUILTIN_STAMPS = [
    {"name": "Approved", "text": "APPROVED", "fg": "#ffffff", "bg": "#10b981"},
    {"name": "Rejected", "text": "REJECTED", "fg": "#ffffff", "bg": "#ef4444"},
    {"name": "Draft", "text": "DRAFT", "fg": "#000000", "bg": "#fbbf24"},
    {"name": "Final", "text": "FINAL", "fg": "#ffffff", "bg": "#2563eb"},
    {"name": "Confidential", "text": "CONFIDENTIAL", "fg": "#ffffff", "bg": "#dc2626"},
    {"name": "For Review", "text": "FOR REVIEW", "fg": "#000000", "bg": "#fb923c"},
    {"name": "Void", "text": "VOID", "fg": "#ffffff", "bg": "#6b7280"},
    {"name": "Copy", "text": "COPY", "fg": "#000000", "bg": "#a3e635"},
]

# ============================================================================
# CONFIGURATION
# ============================================================================

class Config:
    MAX_RECENT_FILES = 15
    MAX_UNDO_STEPS = 100
    DEFAULT_ZOOM = 1.0
    MIN_ZOOM = 0.1
    MAX_ZOOM = 10.0
    
    @staticmethod
    def get_config_path():
        os.makedirs(CONFIG_DIR, exist_ok=True)
        return os.path.join(CONFIG_DIR, "config.json")
    
    @staticmethod
    def load():
        try:
            with open(Config.get_config_path(), 'r') as f:
                return json.load(f)
        except:
            return {"recent_files": [], "window_geometry": "1500x900"}
    
    @staticmethod
    def save(data):
        try:
            with open(Config.get_config_path(), 'w') as f:
                json.dump(data, f, indent=2)
        except:
            pass

# ============================================================================
# ENUMS & DATA CLASSES
# ============================================================================

class ToolMode(Enum):
    SELECT = auto()
    PAN = auto()
    TEXT = auto()
    STICKY_NOTE = auto()
    HIGHLIGHT = auto()
    UNDERLINE = auto()
    STRIKETHROUGH = auto()
    DRAW = auto()
    ERASER = auto()
    RECTANGLE = auto()
    CIRCLE = auto()
    LINE = auto()
    ARROW = auto()
    IMAGE = auto()
    STAMP = auto()
    REDACT = auto()
    CROP = auto()
    LINK = auto()

@dataclass
class SearchResult:
    page: int
    rect: Tuple[float, float, float, float]
    text: str

@dataclass
class Comment:
    id: str
    page: int
    x: float
    y: float
    content: str
    author: str = "User"
    date: str = ""
    color: str = "#fbbf24"

# ============================================================================
# STYLED WIDGETS
# ============================================================================

class ModernButton(tk.Canvas):
    """Modern flat button with hover effects"""
    def __init__(self, parent, text="", icon="", command=None, width=None, 
                 style="default", tooltip="", **kw):
        self.btn_width = width or (36 if not text else max(80, len(text) * 8 + 24))
        self.btn_height = 32
        super().__init__(parent, width=self.btn_width, height=self.btn_height,
                        bg=Theme.BG_SECONDARY, highlightthickness=0, **kw)
        
        self.text = text
        self.icon = icon
        self.command = command
        self.style = style
        self.tooltip_text = tooltip
        self.state = "normal"  # normal, hover, pressed, disabled
        self._tip_window = None
        
        self._draw()
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self.bind("<Button-1>", self._on_press)
        self.bind("<ButtonRelease-1>", self._on_release)
    
    def _get_colors(self):
        if self.state == "disabled":
            return Theme.BG_TERTIARY, Theme.FG_DISABLED
        
        if self.style == "primary":
            if self.state == "pressed":
                return Theme.ACCENT_DARK, Theme.FG_PRIMARY
            elif self.state == "hover":
                return Theme.ACCENT_LIGHT, Theme.FG_PRIMARY
            return Theme.ACCENT, Theme.FG_PRIMARY
        elif self.style == "danger":
            if self.state == "pressed":
                return "#b91c1c", Theme.FG_PRIMARY
            elif self.state == "hover":
                return "#f87171", Theme.FG_PRIMARY
            return Theme.DANGER, Theme.FG_PRIMARY
        else:  # default
            if self.state == "pressed":
                return Theme.BG_ACTIVE, Theme.FG_PRIMARY
            elif self.state == "hover":
                return Theme.BG_HOVER, Theme.FG_PRIMARY
            return Theme.BG_TERTIARY, Theme.FG_SECONDARY
    
    def _draw(self):
        self.delete("all")
        bg, fg = self._get_colors()
        
        # Background with rounded corners effect
        self.create_rectangle(1, 1, self.btn_width-1, self.btn_height-1,
                             fill=bg, outline=Theme.BORDER_LIGHT if self.style == "default" else bg)
        
        # Content
        if self.icon and self.text:
            self.create_text(18, self.btn_height//2, text=self.icon, fill=fg,
                           font=(Theme.FONT_FAMILY, 12))
            self.create_text(36, self.btn_height//2, text=self.text, fill=fg,
                           font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM), anchor="w")
        elif self.icon:
            self.create_text(self.btn_width//2, self.btn_height//2, text=self.icon,
                           fill=fg, font=(Theme.FONT_FAMILY, 14))
        else:
            self.create_text(self.btn_width//2, self.btn_height//2, text=self.text,
                           fill=fg, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
    
    def _on_enter(self, e):
        if self.state != "disabled":
            self.state = "hover"
            self._draw()
            self._show_tooltip()
    
    def _on_leave(self, e):
        if self.state != "disabled":
            self.state = "normal"
            self._draw()
            self._hide_tooltip()
    
    def _on_press(self, e):
        if self.state != "disabled":
            self.state = "pressed"
            self._draw()
    
    def _on_release(self, e):
        if self.state != "disabled":
            self.state = "hover"
            self._draw()
            if self.command and 0 <= e.x <= self.btn_width and 0 <= e.y <= self.btn_height:
                self.command()
    
    def _show_tooltip(self):
        if not self.tooltip_text:
            return
        x = self.winfo_rootx() + self.btn_width // 2
        y = self.winfo_rooty() + self.btn_height + 5
        
        self._tip_window = tk.Toplevel(self)
        self._tip_window.wm_overrideredirect(True)
        self._tip_window.wm_geometry(f"+{x}+{y}")
        
        frame = tk.Frame(self._tip_window, bg=Theme.BG_ELEVATED, padx=8, pady=4)
        frame.pack()
        tk.Label(frame, text=self.tooltip_text, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_XS)).pack()
    
    def _hide_tooltip(self):
        if self._tip_window:
            self._tip_window.destroy()
            self._tip_window = None
    
    def set_state(self, state):
        self.state = state
        self._draw()

class ToolbarButton(tk.Canvas):
    """Toolbar button with icon and optional label"""
    def __init__(self, parent, icon="", label="", command=None, toggle=False, 
                 tooltip="", size="normal", **kw):
        self.size = 48 if size == "normal" else 36
        self.show_label = size == "normal" and label
        height = 56 if self.show_label else self.size
        
        super().__init__(parent, width=self.size, height=height,
                        bg=Theme.BG_SECONDARY, highlightthickness=0, **kw)
        
        self.icon = icon
        self.label = label
        self.command = command
        self.toggle = toggle
        self.tooltip_text = tooltip
        self.active = False
        self.hover = False
        self._tip = None
        
        self._draw()
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self.bind("<Button-1>", self._on_click)
    
    def _draw(self):
        self.delete("all")
        
        # Background
        if self.active:
            self.create_rectangle(2, 2, self.size-2, self.size-2,
                                 fill=Theme.ACCENT_MUTED, outline=Theme.ACCENT)
        elif self.hover:
            self.create_rectangle(2, 2, self.size-2, self.size-2,
                                 fill=Theme.BG_HOVER, outline="")
        
        # Icon
        icon_y = 20 if self.show_label else self.size // 2
        fg = Theme.ACCENT_LIGHT if self.active else (Theme.FG_PRIMARY if self.hover else Theme.FG_SECONDARY)
        self.create_text(self.size//2, icon_y, text=self.icon, fill=fg,
                        font=(Theme.FONT_FAMILY, 16))
        
        # Label
        if self.show_label:
            self.create_text(self.size//2, 42, text=self.label, fill=Theme.FG_MUTED,
                           font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_XS))
    
    def _on_enter(self, e):
        self.hover = True
        self._draw()
        if self.tooltip_text and not self.show_label:
            self._tip = tk.Toplevel(self)
            self._tip.wm_overrideredirect(True)
            self._tip.wm_geometry(f"+{self.winfo_rootx()}+{self.winfo_rooty()+self.size+5}")
            frame = tk.Frame(self._tip, bg=Theme.BG_ELEVATED, padx=6, pady=3)
            frame.pack()
            tk.Label(frame, text=self.tooltip_text, bg=Theme.BG_ELEVATED,
                    fg=Theme.FG_PRIMARY, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_XS)).pack()
    
    def _on_leave(self, e):
        self.hover = False
        self._draw()
        if self._tip:
            self._tip.destroy()
            self._tip = None
    
    def _on_click(self, e):
        if self.toggle:
            self.active = not self.active
        self._draw()
        if self.command:
            self.command()
    
    def set_active(self, active):
        self.active = active
        self._draw()

class ToolbarSeparator(tk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, width=1, height=40, bg=Theme.BORDER_LIGHT, **kw)

class ToolbarGroup(tk.Frame):
    """Group of toolbar buttons with label"""
    def __init__(self, parent, label="", **kw):
        super().__init__(parent, bg=Theme.BG_SECONDARY, **kw)
        
        self.buttons_frame = tk.Frame(self, bg=Theme.BG_SECONDARY)
        self.buttons_frame.pack(pady=(4, 2))
        
        if label:
            tk.Label(self, text=label, bg=Theme.BG_SECONDARY, fg=Theme.FG_MUTED,
                    font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_XS)).pack()
    
    def add_button(self, **kw):
        btn = ToolbarButton(self.buttons_frame, **kw)
        btn.pack(side=tk.LEFT, padx=1)
        return btn

class ModernEntry(tk.Entry):
    """Styled entry widget"""
    def __init__(self, parent, placeholder="", **kw):
        super().__init__(parent, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY,
                        insertbackground=Theme.FG_PRIMARY, relief=tk.FLAT,
                        font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM),
                        highlightthickness=1, highlightcolor=Theme.BORDER_FOCUS,
                        highlightbackground=Theme.BORDER_LIGHT, **kw)
        
        self.placeholder = placeholder
        self._has_placeholder = False
        
        if placeholder:
            self._show_placeholder()
            self.bind("<FocusIn>", self._on_focus_in)
            self.bind("<FocusOut>", self._on_focus_out)
    
    def _show_placeholder(self):
        if not self.get():
            self._has_placeholder = True
            self.insert(0, self.placeholder)
            self.configure(fg=Theme.FG_MUTED)
    
    def _on_focus_in(self, e):
        if self._has_placeholder:
            self.delete(0, tk.END)
            self.configure(fg=Theme.FG_PRIMARY)
            self._has_placeholder = False
    
    def _on_focus_out(self, e):
        if not self.get():
            self._show_placeholder()
    
    def get_value(self):
        if self._has_placeholder:
            return ""
        return self.get()

class TabButton(tk.Canvas):
    """Document tab button"""
    def __init__(self, parent, title="", doc_id="", on_select=None, on_close=None, **kw):
        super().__init__(parent, width=180, height=Theme.TAB_HEIGHT,
                        bg=Theme.BG_PRIMARY, highlightthickness=0, **kw)
        
        self.title = title
        self.doc_id = doc_id
        self.on_select = on_select
        self.on_close = on_close
        self.active = False
        self.hover = False
        self.close_hover = False
        
        self._draw()
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self.bind("<Button-1>", self._on_click)
        self.bind("<Motion>", self._on_motion)
    
    def _draw(self):
        self.delete("all")
        
        # Background
        bg = Theme.BG_TERTIARY if self.active else (Theme.BG_SECONDARY if self.hover else Theme.BG_PRIMARY)
        self.create_rectangle(0, 0, 180, Theme.TAB_HEIGHT, fill=bg, outline="")
        
        # Active indicator
        if self.active:
            self.create_rectangle(0, Theme.TAB_HEIGHT - 2, 180, Theme.TAB_HEIGHT,
                                 fill=Theme.ACCENT, outline="")
        
        # Icon
        self.create_text(16, Theme.TAB_HEIGHT//2, text="📄", font=(Theme.FONT_FAMILY, 10))
        
        # Title
        display_title = self.title[:18] + "..." if len(self.title) > 18 else self.title
        self.create_text(30, Theme.TAB_HEIGHT//2, text=display_title,
                        fill=Theme.FG_PRIMARY if self.active else Theme.FG_SECONDARY,
                        font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM), anchor="w")
        
        # Close button
        close_bg = Theme.BG_HOVER if self.close_hover else ""
        if close_bg:
            self.create_oval(152, 8, 172, 28, fill=close_bg, outline="")
        self.create_text(162, Theme.TAB_HEIGHT//2, text="×",
                        fill=Theme.FG_PRIMARY if self.close_hover else Theme.FG_MUTED,
                        font=(Theme.FONT_FAMILY, 14))
    
    def _on_enter(self, e):
        self.hover = True
        self._draw()
    
    def _on_leave(self, e):
        self.hover = False
        self.close_hover = False
        self._draw()
    
    def _on_motion(self, e):
        in_close = 152 <= e.x <= 172 and 8 <= e.y <= 28
        if in_close != self.close_hover:
            self.close_hover = in_close
            self._draw()
    
    def _on_click(self, e):
        if 152 <= e.x <= 172 and 8 <= e.y <= 28:
            if self.on_close:
                self.on_close(self.doc_id)
        else:
            if self.on_select:
                self.on_select(self.doc_id)
    
    def set_active(self, active):
        self.active = active
        self._draw()
    
    def set_title(self, title):
        self.title = title
        self._draw()

class SidebarTab(tk.Canvas):
    """Sidebar navigation tab"""
    def __init__(self, parent, icon="", label="", command=None, **kw):
        super().__init__(parent, width=Theme.SIDEBAR_WIDTH, height=40,
                        bg=Theme.BG_SECONDARY, highlightthickness=0, **kw)
        
        self.icon = icon
        self.label = label
        self.command = command
        self.active = False
        self.hover = False
        
        self._draw()
        self.bind("<Enter>", lambda e: self._set_hover(True))
        self.bind("<Leave>", lambda e: self._set_hover(False))
        self.bind("<Button-1>", self._on_click)
    
    def _draw(self):
        self.delete("all")
        
        if self.active:
            self.create_rectangle(0, 0, 3, 40, fill=Theme.ACCENT, outline="")
            self.create_rectangle(3, 0, Theme.SIDEBAR_WIDTH, 40, fill=Theme.BG_TERTIARY, outline="")
            fg = Theme.FG_PRIMARY
        elif self.hover:
            self.create_rectangle(0, 0, Theme.SIDEBAR_WIDTH, 40, fill=Theme.BG_HOVER, outline="")
            fg = Theme.FG_PRIMARY
        else:
            fg = Theme.FG_SECONDARY
        
        self.create_text(24, 20, text=self.icon, fill=fg, font=(Theme.FONT_FAMILY, 14))
        self.create_text(48, 20, text=self.label, fill=fg,
                        font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM), anchor="w")
    
    def _set_hover(self, h):
        self.hover = h
        self._draw()
    
    def _on_click(self, e):
        if self.command:
            self.command()
    
    def set_active(self, active):
        self.active = active
        self._draw()

# ============================================================================
# PDF DOCUMENT CLASS
# ============================================================================

class PDFDocument:
    def __init__(self):
        self.doc = None
        self.filepath = None
        self.is_modified = False
        self.comments = []
        self._comment_counter = 0
    
    def open(self, filepath):
        try:
            self.doc = fitz.open(filepath)
            self.filepath = filepath
            self.is_modified = False
            self.comments = []
            self._load_comments()
            return True
        except Exception as e:
            print(f"Open error: {e}")
            return False
    
    def create_new(self, width=612, height=792):
        self.doc = fitz.open()
        self.doc.new_page(width=width, height=height)
        self.filepath = None
        self.is_modified = True
    
    def save(self, filepath=None):
        if not self.doc:
            return False
        path = filepath or self.filepath
        if not path:
            return False
        try:
            self._save_comments()
            if path == self.filepath:
                self.doc.saveIncr()
            else:
                self.doc.save(path, garbage=4, deflate=True)
            self.filepath = path
            self.is_modified = False
            return True
        except:
            return False
    
    def close(self):
        if self.doc:
            self.doc.close()
        self.__init__()
    
    @property
    def page_count(self):
        return len(self.doc) if self.doc else 0
    
    @property
    def filename(self):
        return os.path.basename(self.filepath) if self.filepath else "Untitled"
    
    def get_page(self, num):
        if self.doc and 0 <= num < len(self.doc):
            return self.doc[num]
        return None
    
    def render_page(self, page_num, zoom=1.0):
        page = self.get_page(page_num)
        if not page:
            return None
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    
    def get_page_size(self, page_num):
        page = self.get_page(page_num)
        return (page.rect.width, page.rect.height) if page else (612, 792)
    
    def get_text(self, page_num):
        page = self.get_page(page_num)
        return page.get_text() if page else ""
    
    def search_text(self, query, case_sensitive=False):
        results = []
        if not self.doc or not query:
            return results
        for i in range(len(self.doc)):
            for rect in self.doc[i].search_for(query):
                results.append(SearchResult(i, tuple(rect), query))
        return results
    
    def delete_page(self, page_num):
        if self.doc and 0 <= page_num < len(self.doc) and len(self.doc) > 1:
            self.doc.delete_page(page_num)
            self.is_modified = True
            return True
        return False
    
    def insert_page(self, index=-1, width=612, height=792):
        if self.doc:
            if index < 0:
                index = len(self.doc)
            self.doc.new_page(pno=index, width=width, height=height)
            self.is_modified = True
    
    def duplicate_page(self, page_num):
        if self.doc and 0 <= page_num < len(self.doc):
            self.doc.fullcopy_page(page_num, page_num + 1)
            self.is_modified = True
    
    def rotate_page(self, page_num, angle=90):
        page = self.get_page(page_num)
        if page:
            page.set_rotation((page.rotation + angle) % 360)
            self.is_modified = True
    
    def crop_page(self, page_num, rect):
        page = self.get_page(page_num)
        if page:
            page.set_cropbox(fitz.Rect(rect))
            self.is_modified = True
    
    # Annotations
    def add_text(self, page_num, text, x, y, font_size=12, color=(0, 0, 0)):
        page = self.get_page(page_num)
        if page and text:
            fitz_color = tuple(c/255 for c in color) if max(color) > 1 else color
            page.insert_text((x, y), text, fontsize=font_size, fontname="helv", color=fitz_color)
            self.is_modified = True
    
    def add_highlight(self, page_num, rect, color=(1, 1, 0)):
        page = self.get_page(page_num)
        if page:
            annot = page.add_highlight_annot(fitz.Rect(rect))
            annot.set_colors(stroke=color)
            annot.update()
            self.is_modified = True
    
    def add_underline(self, page_num, rect):
        page = self.get_page(page_num)
        if page:
            page.add_underline_annot(fitz.Rect(rect)).update()
            self.is_modified = True
    
    def add_strikethrough(self, page_num, rect):
        page = self.get_page(page_num)
        if page:
            page.add_strikeout_annot(fitz.Rect(rect)).update()
            self.is_modified = True
    
    def add_rect(self, page_num, rect, color=(1, 0, 0), width=2):
        page = self.get_page(page_num)
        if page:
            shape = page.new_shape()
            shape.draw_rect(fitz.Rect(rect))
            shape.finish(color=color, width=width)
            shape.commit()
            self.is_modified = True
    
    def add_circle(self, page_num, rect, color=(1, 0, 0), width=2):
        page = self.get_page(page_num)
        if page:
            shape = page.new_shape()
            shape.draw_oval(fitz.Rect(rect))
            shape.finish(color=color, width=width)
            shape.commit()
            self.is_modified = True
    
    def add_line(self, page_num, p1, p2, color=(0, 0, 0), width=2):
        page = self.get_page(page_num)
        if page:
            shape = page.new_shape()
            shape.draw_line(p1, p2)
            shape.finish(color=color, width=width)
            shape.commit()
            self.is_modified = True
    
    def add_arrow(self, page_num, p1, p2, color=(0, 0, 0)):
        page = self.get_page(page_num)
        if page:
            annot = page.add_line_annot(fitz.Point(p1), fitz.Point(p2))
            annot.set_colors(stroke=color)
            annot.set_line_ends(fitz.PDF_ANNOT_LE_NONE, fitz.PDF_ANNOT_LE_CLOSED_ARROW)
            annot.set_border(width=2)
            annot.update()
            self.is_modified = True
    
    def add_freehand(self, page_num, points, color=(0, 0, 0), width=2):
        page = self.get_page(page_num)
        if page and len(points) >= 2:
            annot = page.add_ink_annot([points])
            annot.set_colors(stroke=color)
            annot.set_border(width=width)
            annot.update()
            self.is_modified = True
    
    def add_image(self, page_num, image_path, x=None, y=None, width=None, height=None):
        page = self.get_page(page_num)
        if not page:
            return False
        try:
            img = Image.open(image_path)
            iw, ih = img.size
            if width and not height:
                height = width * ih / iw
            elif height and not width:
                width = height * iw / ih
            elif not width and not height:
                scale = min(300/iw, 300/ih, 1.0)
                width, height = iw * scale, ih * scale
            if x is None:
                x = (page.rect.width - width) / 2
            if y is None:
                y = (page.rect.height - height) / 2
            page.insert_image(fitz.Rect(x, y, x+width, y+height), filename=image_path)
            self.is_modified = True
            return True
        except:
            return False
    
    def add_stamp(self, page_num, x, y, stamp):
        page = self.get_page(page_num)
        if not page:
            return
        text = stamp['text']
        font_size = 14
        text_width = len(text) * font_size * 0.6
        stamp_w, stamp_h = text_width + 20, font_size + 16
        
        def hex_to_rgb(h):
            h = h.lstrip('#')
            return tuple(int(h[i:i+2], 16)/255 for i in (0, 2, 4))
        
        bg = hex_to_rgb(stamp['bg'])
        fg = hex_to_rgb(stamp['fg'])
        
        shape = page.new_shape()
        shape.draw_rect(fitz.Rect(x, y, x + stamp_w, y + stamp_h))
        shape.finish(color=bg, fill=bg, width=2)
        shape.commit()
        
        page.insert_text((x + 10, y + stamp_h - 8), text, fontsize=font_size, fontname="hebo", color=fg)
        self.is_modified = True
    
    def redact_area(self, page_num, rect):
        page = self.get_page(page_num)
        if page:
            page.add_redact_annot(fitz.Rect(rect), fill=(0, 0, 0))
            page.apply_redactions()
            self.is_modified = True
    
    # Comments
    def add_comment(self, page, x, y, content, author="User"):
        self._comment_counter += 1
        comment = Comment(f"c_{self._comment_counter}", page, x, y, content, author,
                         datetime.now().strftime("%Y-%m-%d %H:%M"))
        self.comments.append(comment)
        self.is_modified = True
        return comment
    
    def _load_comments(self):
        if not self.doc:
            return
        for i, page in enumerate(self.doc):
            for annot in page.annots():
                if annot.type[0] == fitz.PDF_ANNOT_TEXT:
                    self._comment_counter += 1
                    rect = annot.rect
                    self.comments.append(Comment(
                        f"c_{self._comment_counter}", i, rect.x0, rect.y0,
                        annot.info.get("content", ""),
                        annot.info.get("title", "User")
                    ))
    
    def _save_comments(self):
        if not self.doc:
            return
        for page in self.doc:
            for a in [ann for ann in page.annots() if ann.type[0] == fitz.PDF_ANNOT_TEXT]:
                page.delete_annot(a)
        for c in self.comments:
            page = self.get_page(c.page)
            if page:
                annot = page.add_text_annot((c.x, c.y), c.content)
                annot.set_info(title=c.author)
                annot.update()
    
    # Bookmarks
    def get_bookmarks(self):
        if not self.doc:
            return []
        return [(item[0], item[1], item[2]-1) for item in self.doc.get_toc()]
    
    # Form fields
    def get_form_fields(self, page_num=None):
        fields = []
        if not self.doc:
            return fields
        pages = [page_num] if page_num is not None else range(len(self.doc))
        for pnum in pages:
            page = self.get_page(pnum)
            if page:
                for widget in page.widgets():
                    fields.append({
                        'page': pnum, 'name': widget.field_name,
                        'type': widget.field_type_string,
                        'value': widget.field_value or '',
                        'rect': tuple(widget.rect), 'widget': widget
                    })
        return fields
    
    def set_form_field(self, page_num, name, value):
        page = self.get_page(page_num)
        if page:
            for widget in page.widgets():
                if widget.field_name == name:
                    widget.field_value = value
                    widget.update()
                    self.is_modified = True
                    return True
        return False
    
    # Document operations
    def compress(self, output_path):
        if self.doc:
            try:
                self.doc.save(output_path, garbage=4, deflate=True, clean=True, linear=True)
                return True
            except:
                pass
        return False
    
    def add_watermark(self, text, font_size=48, color=(0.8, 0.8, 0.8), angle=45):
        if not self.doc:
            return
        for page in self.doc:
            rect = page.rect
            cx, cy = rect.width / 2, rect.height / 2
            text_width = len(text) * font_size * 0.5
            page.insert_text(fitz.Point(cx - text_width/2, cy), text,
                           fontsize=font_size, fontname="helv", color=color, rotate=angle)
        self.is_modified = True
    
    def add_header_footer(self, header=None, footer=None, font_size=10, margin=36):
        if not self.doc:
            return
        for i, page in enumerate(self.doc):
            pw, ph = page.rect.width, page.rect.height
            page_num = i + 1
            
            def process(txt):
                if not txt:
                    return None
                return txt.replace("{page}", str(page_num)).replace("{pages}", str(len(self.doc))).replace("{date}", datetime.now().strftime("%Y-%m-%d"))
            
            if header:
                h = process(header)
                x = (pw - len(h) * font_size * 0.4) / 2
                page.insert_text((x, margin), h, fontsize=font_size, fontname="helv", color=(0, 0, 0))
            if footer:
                f = process(footer)
                x = (pw - len(f) * font_size * 0.4) / 2
                page.insert_text((x, ph - margin + font_size), f, fontsize=font_size, fontname="helv", color=(0, 0, 0))
        self.is_modified = True
    
    def add_bates_numbers(self, prefix="", start=1, digits=6, position="bottom-right", font_size=10, margin=36):
        if not self.doc:
            return
        for i, page in enumerate(self.doc):
            bates = f"{prefix}{start + i:0{digits}d}"
            pw, ph = page.rect.width, page.rect.height
            tw = len(bates) * font_size * 0.5
            positions = {
                "top-left": (margin, margin + font_size),
                "top-right": (pw - tw - margin, margin + font_size),
                "bottom-left": (margin, ph - margin),
                "bottom-right": (pw - tw - margin, ph - margin),
            }
            x, y = positions.get(position, positions["bottom-right"])
            page.insert_text((x, y), bates, fontsize=font_size, fontname="helv", color=(0, 0, 0))
        self.is_modified = True
    
    def flatten_annotations(self):
        if self.doc:
            for page in self.doc:
                page.clean_contents()
            self.is_modified = True
    
    def remove_metadata(self):
        if self.doc:
            self.doc.set_metadata({})
            self.is_modified = True
    
    def get_metadata(self):
        return dict(self.doc.metadata) if self.doc else {}
    
    def set_metadata(self, data):
        if self.doc:
            self.doc.set_metadata(data)
            self.is_modified = True
    
    def export_to_word(self, output_path):
        if not HAS_DOCX or not self.doc:
            return False
        try:
            doc = DocxDocument()
            for i in range(len(self.doc)):
                if i > 0:
                    doc.add_page_break()
                text = self.get_text(i)
                if text.strip():
                    doc.add_paragraph(text)
            doc.save(output_path)
            return True
        except:
            return False
    
    def export_to_images(self, output_dir, dpi=150, fmt="png"):
        files = []
        if not self.doc:
            return files
        zoom = dpi / 72
        for i in range(len(self.doc)):
            pix = self.doc[i].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            path = os.path.join(output_dir, f"page_{i+1:03d}.{fmt}")
            pix.save(path)
            files.append(path)
        return files
    
    def export_text(self, output_path):
        if not self.doc:
            return False
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                for i in range(len(self.doc)):
                    f.write(f"--- Page {i+1} ---\n{self.get_text(i)}\n\n")
            return True
        except:
            return False
    
    def merge_pdf(self, other_path):
        if self.doc:
            other = fitz.open(other_path)
            self.doc.insert_pdf(other)
            other.close()
            self.is_modified = True
    
    def split_pages(self, output_dir):
        files = []
        if not self.doc:
            return files
        for i in range(len(self.doc)):
            new_doc = fitz.open()
            new_doc.insert_pdf(self.doc, from_page=i, to_page=i)
            path = os.path.join(output_dir, f"page_{i+1:03d}.pdf")
            new_doc.save(path)
            new_doc.close()
            files.append(path)
        return files

# ============================================================================
# OCR ENGINE
# ============================================================================

class OCREngine:
    @staticmethod
    def is_available():
        try:
            import pytesseract
            OCREngine._configure()
            pytesseract.get_tesseract_version()
            return True, "OK"
        except ImportError:
            return False, "pytesseract not installed"
        except:
            return False, "Tesseract not found"
    
    @staticmethod
    def _configure():
        try:
            import pytesseract
            if path := os.environ.get("TESSERACT_CMD"):
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
        except:
            pass
    
    @staticmethod
    def make_searchable(doc, callback=None):
        try:
            import pytesseract
            OCREngine._configure()
        except:
            return False, 0
        
        processed = 0
        for pnum in range(doc.page_count):
            page = doc.get_page(pnum)
            if not page:
                continue
            if callback:
                callback(f"Processing page {pnum + 1}...")
            
            img = doc.render_page(pnum, zoom=2.0)
            if not img:
                continue
            
            pw, ph = page.rect.width, page.rect.height
            iw, ih = img.size
            sx, sy = pw / iw, ph / ih
            
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            
            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                conf = int(data['conf'][i]) if str(data['conf'][i]).lstrip('-').isdigit() else 0
                if not text or conf < 30:
                    continue
                
                x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                px, py, pw_t, ph_t = x * sx, y * sy, w * sx, h * sy
                
                fs = max(4, min(72, ph_t * 0.85))
                try:
                    tl = fitz.get_text_length(text, fontsize=fs)
                    if tl > 0 and pw_t > 0:
                        fs = max(4, min(72, fs * (pw_t / tl)))
                    page.insert_text((px, py + ph_t * 0.85), text, fontsize=fs,
                                    fontname="helv", color=(0, 0, 0), render_mode=3)
                except:
                    pass
            processed += 1
        
        if processed > 0:
            doc.is_modified = True
        return processed > 0, processed

# ============================================================================
# MAIN APPLICATION
# ============================================================================

class PDFEditorPro(tk.Tk):
    def __init__(self):
        super().__init__()
        
        self.title("PDF Editor Pro")
        self.geometry("1500x900")
        self.minsize(1200, 750)
        self.configure(bg=Theme.BG_DARK)
        
        # State
        self.documents = {}
        self.active_doc_id = None
        self.current_page = 0
        self.zoom = 1.0
        self.tool_mode = ToolMode.SELECT
        self.draw_color = (0, 0, 0)
        self.draw_points = []
        self.drag_start = None
        self.page_image = None
        self.search_results = []
        self.selected_stamp = None
        self.sidebar_mode = "pages"
        
        self.config_data = Config.load()
        
        self._build_menu()
        self._build_ui()
        self._bind_shortcuts()
        
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self._show_welcome()
    
    @property
    def doc(self):
        return self.documents.get(self.active_doc_id)
    
    # =========================================================================
    # UI BUILDING
    # =========================================================================
    
    def _build_menu(self):
        menubar = tk.Menu(self, bg=Theme.BG_TERTIARY, fg=Theme.FG_PRIMARY,
                         activebackground=Theme.ACCENT, activeforeground=Theme.FG_PRIMARY,
                         font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        
        # File
        file_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                           activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        file_menu.add_command(label="New", command=self._new_doc, accelerator="Ctrl+N")
        file_menu.add_command(label="Open...", command=self._open_doc, accelerator="Ctrl+O")
        file_menu.add_separator()
        file_menu.add_command(label="Save", command=self._save_doc, accelerator="Ctrl+S")
        file_menu.add_command(label="Save As...", command=self._save_as)
        file_menu.add_separator()
        file_menu.add_command(label="Close", command=self._close_tab, accelerator="Ctrl+W")
        file_menu.add_separator()
        file_menu.add_command(label="Properties...", command=self._show_properties)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self._on_close)
        menubar.add_cascade(label="File", menu=file_menu)
        
        # Edit
        edit_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                           activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        edit_menu.add_command(label="Find...", command=self._show_search, accelerator="Ctrl+F")
        edit_menu.add_separator()
        edit_menu.add_command(label="Copy Text", command=self._copy_text)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        
        # View
        view_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                           activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        view_menu.add_command(label="Zoom In", command=self._zoom_in, accelerator="Ctrl++")
        view_menu.add_command(label="Zoom Out", command=self._zoom_out, accelerator="Ctrl+-")
        view_menu.add_command(label="Fit Page", command=self._zoom_fit)
        view_menu.add_command(label="Actual Size", command=self._zoom_100)
        view_menu.add_separator()
        view_menu.add_command(label="Rotate CW", command=lambda: self._rotate(90))
        view_menu.add_command(label="Rotate CCW", command=lambda: self._rotate(-90))
        menubar.add_cascade(label="View", menu=view_menu)
        
        # Page
        page_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                           activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        page_menu.add_command(label="Insert Page", command=self._insert_page)
        page_menu.add_command(label="Duplicate Page", command=self._duplicate_page)
        page_menu.add_command(label="Delete Page", command=self._delete_page)
        page_menu.add_separator()
        page_menu.add_command(label="Extract Page...", command=self._extract_page)
        menubar.add_cascade(label="Page", menu=page_menu)
        
        # Tools
        tools_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                            activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        tools_menu.add_command(label="Add Text", command=lambda: self._set_tool(ToolMode.TEXT))
        tools_menu.add_command(label="Add Comment", command=lambda: self._set_tool(ToolMode.STICKY_NOTE))
        tools_menu.add_command(label="Add Image...", command=self._add_image)
        tools_menu.add_command(label="Add Stamp...", command=self._show_stamp_dialog)
        tools_menu.add_separator()
        tools_menu.add_command(label="Highlight", command=lambda: self._set_tool(ToolMode.HIGHLIGHT))
        tools_menu.add_command(label="Underline", command=lambda: self._set_tool(ToolMode.UNDERLINE))
        tools_menu.add_command(label="Strikethrough", command=lambda: self._set_tool(ToolMode.STRIKETHROUGH))
        tools_menu.add_separator()
        tools_menu.add_command(label="Redact", command=lambda: self._set_tool(ToolMode.REDACT))
        menubar.add_cascade(label="Tools", menu=tools_menu)
        
        # Document
        doc_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                          activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        doc_menu.add_command(label="Merge PDFs...", command=self._merge_pdfs)
        doc_menu.add_command(label="Split Document...", command=self._split_doc)
        doc_menu.add_separator()
        doc_menu.add_command(label="Add Watermark...", command=self._watermark_dialog)
        doc_menu.add_command(label="Add Header/Footer...", command=self._header_footer_dialog)
        doc_menu.add_command(label="Bates Numbering...", command=self._bates_dialog)
        doc_menu.add_separator()
        doc_menu.add_command(label="OCR - Make Searchable", command=self._ocr_doc)
        doc_menu.add_separator()
        doc_menu.add_command(label="Compress...", command=self._compress_doc)
        doc_menu.add_command(label="Password Protect...", command=self._password_dialog)
        menubar.add_cascade(label="Document", menu=doc_menu)
        
        # Export
        export_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                             activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        export_menu.add_command(label="Export to Word...", command=self._export_word)
        export_menu.add_command(label="Export to Images...", command=self._export_images)
        export_menu.add_command(label="Export Text...", command=self._export_text)
        menubar.add_cascade(label="Export", menu=export_menu)
        
        # Help
        help_menu = tk.Menu(menubar, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY,
                           activebackground=Theme.ACCENT, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        help_menu.add_command(label="Keyboard Shortcuts", command=self._show_shortcuts)
        help_menu.add_separator()
        help_menu.add_command(label="About", command=self._show_about)
        menubar.add_cascade(label="Help", menu=help_menu)
        
        self.config(menu=menubar)
    
    def _build_ui(self):
        # Main container
        main = tk.Frame(self, bg=Theme.BG_DARK)
        main.pack(fill=tk.BOTH, expand=True)
        
        # Tab bar
        self.tab_bar = tk.Frame(main, bg=Theme.BG_PRIMARY, height=Theme.TAB_HEIGHT)
        self.tab_bar.pack(fill=tk.X)
        self.tab_bar.pack_propagate(False)
        self.tabs = {}
        
        # New tab button
        new_tab_btn = tk.Label(self.tab_bar, text=" + ", bg=Theme.BG_PRIMARY, fg=Theme.FG_MUTED,
                              font=(Theme.FONT_FAMILY, 14), cursor="hand2")
        new_tab_btn.pack(side=tk.LEFT, padx=5, pady=5)
        new_tab_btn.bind("<Button-1>", lambda e: self._new_doc())
        new_tab_btn.bind("<Enter>", lambda e: new_tab_btn.configure(fg=Theme.FG_PRIMARY))
        new_tab_btn.bind("<Leave>", lambda e: new_tab_btn.configure(fg=Theme.FG_MUTED))
        
        # Toolbar
        self._build_toolbar(main)
        
        # Content area
        content = tk.Frame(main, bg=Theme.BG_DARK)
        content.pack(fill=tk.BOTH, expand=True)
        
        # Left sidebar
        self._build_sidebar(content)
        
        # Canvas area
        self._build_canvas(content)
        
        # Right panel (properties)
        self._build_properties_panel(content)
        
        # Status bar
        self._build_status_bar(main)
    
    def _build_toolbar(self, parent):
        toolbar = tk.Frame(parent, bg=Theme.BG_SECONDARY, height=Theme.TOOLBAR_HEIGHT)
        toolbar.pack(fill=tk.X)
        toolbar.pack_propagate(False)
        
        # File group
        file_group = ToolbarGroup(toolbar, label="File")
        file_group.pack(side=tk.LEFT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
        file_group.add_button(icon="📄", label="New", command=self._new_doc, tooltip="New Document")
        file_group.add_button(icon="📂", label="Open", command=self._open_doc, tooltip="Open File")
        file_group.add_button(icon="💾", label="Save", command=self._save_doc, tooltip="Save")
        
        ToolbarSeparator(toolbar).pack(side=tk.LEFT, padx=Theme.PAD_SM, pady=Theme.PAD_LG)
        
        # Tools group
        tools_group = ToolbarGroup(toolbar, label="Tools")
        tools_group.pack(side=tk.LEFT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
        
        self.tool_buttons = {}
        tools = [
            (ToolMode.SELECT, "👆", "Select"),
            (ToolMode.PAN, "✋", "Pan"),
            (ToolMode.TEXT, "T", "Text"),
            (ToolMode.STICKY_NOTE, "📝", "Comment"),
            (ToolMode.HIGHLIGHT, "🔆", "Highlight"),
            (ToolMode.DRAW, "✏️", "Draw"),
        ]
        for mode, icon, label in tools:
            btn = tools_group.add_button(icon=icon, label=label, command=lambda m=mode: self._set_tool(m),
                                        toggle=True, tooltip=label)
            self.tool_buttons[mode] = btn
        self.tool_buttons[ToolMode.SELECT].set_active(True)
        
        ToolbarSeparator(toolbar).pack(side=tk.LEFT, padx=Theme.PAD_SM, pady=Theme.PAD_LG)
        
        # Shapes group
        shapes_group = ToolbarGroup(toolbar, label="Shapes")
        shapes_group.pack(side=tk.LEFT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
        shapes = [
            (ToolMode.RECTANGLE, "▢", "Rectangle"),
            (ToolMode.CIRCLE, "○", "Circle"),
            (ToolMode.ARROW, "↗", "Arrow"),
            (ToolMode.LINE, "╱", "Line"),
        ]
        for mode, icon, label in shapes:
            btn = shapes_group.add_button(icon=icon, label=label, command=lambda m=mode: self._set_tool(m),
                                         toggle=True, tooltip=label)
            self.tool_buttons[mode] = btn
        
        ToolbarSeparator(toolbar).pack(side=tk.LEFT, padx=Theme.PAD_SM, pady=Theme.PAD_LG)
        
        # Insert group
        insert_group = ToolbarGroup(toolbar, label="Insert")
        insert_group.pack(side=tk.LEFT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
        insert_group.add_button(icon="🖼", label="Image", command=self._add_image, tooltip="Insert Image")
        insert_group.add_button(icon="📌", label="Stamp", command=self._show_stamp_dialog, tooltip="Add Stamp")
        btn = insert_group.add_button(icon="▮", label="Redact", command=lambda: self._set_tool(ToolMode.REDACT),
                                     toggle=True, tooltip="Redact Area")
        self.tool_buttons[ToolMode.REDACT] = btn
        
        # Right side - Navigation
        nav_frame = tk.Frame(toolbar, bg=Theme.BG_SECONDARY)
        nav_frame.pack(side=tk.RIGHT, padx=Theme.PAD_LG, pady=Theme.PAD_MD)
        
        # Zoom controls
        zoom_frame = tk.Frame(nav_frame, bg=Theme.BG_SECONDARY)
        zoom_frame.pack(side=tk.RIGHT, padx=(Theme.PAD_LG, 0))
        
        ToolbarButton(zoom_frame, icon="−", command=self._zoom_out, tooltip="Zoom Out", size="small").pack(side=tk.LEFT)
        self.zoom_label = tk.Label(zoom_frame, text="100%", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY,
                                   width=6, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.zoom_label.pack(side=tk.LEFT, padx=Theme.PAD_SM)
        ToolbarButton(zoom_frame, icon="+", command=self._zoom_in, tooltip="Zoom In", size="small").pack(side=tk.LEFT)
        ToolbarButton(zoom_frame, icon="⊡", command=self._zoom_fit, tooltip="Fit Page", size="small").pack(side=tk.LEFT, padx=(Theme.PAD_SM, 0))
        
        # Page navigation
        page_frame = tk.Frame(nav_frame, bg=Theme.BG_SECONDARY)
        page_frame.pack(side=tk.RIGHT, padx=Theme.PAD_LG)
        
        ToolbarButton(page_frame, icon="⏮", command=self._first_page, tooltip="First Page", size="small").pack(side=tk.LEFT)
        ToolbarButton(page_frame, icon="◀", command=self._prev_page, tooltip="Previous", size="small").pack(side=tk.LEFT)
        
        self.page_entry = ModernEntry(page_frame, width=5)
        self.page_entry.pack(side=tk.LEFT, padx=Theme.PAD_SM, ipady=2)
        self.page_entry.bind("<Return>", self._goto_page_entry)
        
        self.page_total = tk.Label(page_frame, text="/ 0", bg=Theme.BG_SECONDARY, fg=Theme.FG_SECONDARY,
                                   font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.page_total.pack(side=tk.LEFT)
        
        ToolbarButton(page_frame, icon="▶", command=self._next_page, tooltip="Next", size="small").pack(side=tk.LEFT, padx=(Theme.PAD_SM, 0))
        ToolbarButton(page_frame, icon="⏭", command=self._last_page, tooltip="Last Page", size="small").pack(side=tk.LEFT)
    
    def _build_sidebar(self, parent):
        self.sidebar = tk.Frame(parent, bg=Theme.BG_SECONDARY, width=Theme.SIDEBAR_WIDTH)
        self.sidebar.pack(side=tk.LEFT, fill=tk.Y)
        self.sidebar.pack_propagate(False)
        
        # Sidebar tabs
        self.sidebar_tabs = {}
        tabs_data = [
            ("pages", "📄", "Pages"),
            ("bookmarks", "📑", "Bookmarks"),
            ("comments", "💬", "Comments"),
        ]
        for key, icon, label in tabs_data:
            tab = SidebarTab(self.sidebar, icon=icon, label=label,
                            command=lambda k=key: self._show_sidebar_content(k))
            tab.pack(fill=tk.X)
            self.sidebar_tabs[key] = tab
        
        self.sidebar_tabs["pages"].set_active(True)
        
        # Separator
        tk.Frame(self.sidebar, height=1, bg=Theme.BORDER_LIGHT).pack(fill=tk.X, pady=Theme.PAD_SM)
        
        # Content area
        self.sidebar_content = tk.Frame(self.sidebar, bg=Theme.BG_SECONDARY)
        self.sidebar_content.pack(fill=tk.BOTH, expand=True)
        
        # Pages panel
        self._build_pages_panel()
        
        # Bookmarks panel
        self.bookmarks_panel = tk.Frame(self.sidebar_content, bg=Theme.BG_SECONDARY)
        self.bookmarks_tree = ttk.Treeview(self.bookmarks_panel, show="tree")
        self.bookmarks_tree.pack(fill=tk.BOTH, expand=True, padx=Theme.PAD_SM, pady=Theme.PAD_SM)
        self.bookmarks_tree.bind("<<TreeviewSelect>>", self._on_bookmark_select)
        
        # Comments panel
        self.comments_panel = tk.Frame(self.sidebar_content, bg=Theme.BG_SECONDARY)
        self.comments_list = tk.Listbox(self.comments_panel, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY,
                                        selectbackground=Theme.ACCENT, relief=tk.FLAT,
                                        font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.comments_list.pack(fill=tk.BOTH, expand=True, padx=Theme.PAD_SM, pady=Theme.PAD_SM)
        self.comments_list.bind("<<ListboxSelect>>", self._on_comment_select)
        
        self._show_sidebar_content("pages")
    
    def _build_pages_panel(self):
        self.pages_panel = tk.Frame(self.sidebar_content, bg=Theme.BG_SECONDARY)
        
        # Scrollable thumbnail area
        canvas = tk.Canvas(self.pages_panel, bg=Theme.BG_SECONDARY, highlightthickness=0, width=180)
        scrollbar = ttk.Scrollbar(self.pages_panel, orient=tk.VERTICAL, command=canvas.yview)
        self.thumb_frame = tk.Frame(canvas, bg=Theme.BG_SECONDARY)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.thumb_canvas_window = canvas.create_window((0, 0), window=self.thumb_frame, anchor=tk.NW)
        self.thumb_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<MouseWheel>", lambda e: canvas.yview_scroll(-1 * (e.delta // 120), "units"))
        
        self.thumbnails = []
    
    def _build_canvas(self, parent):
        canvas_container = tk.Frame(parent, bg=Theme.BG_DARK)
        canvas_container.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Search bar (hidden initially)
        self.search_frame = tk.Frame(canvas_container, bg=Theme.BG_TERTIARY, height=44)
        
        tk.Label(self.search_frame, text="🔍", bg=Theme.BG_TERTIARY, fg=Theme.FG_MUTED).pack(side=tk.LEFT, padx=(Theme.PAD_MD, Theme.PAD_SM))
        self.search_entry = ModernEntry(self.search_frame, width=30, placeholder="Find in document...")
        self.search_entry.pack(side=tk.LEFT, padx=Theme.PAD_SM, pady=Theme.PAD_SM, ipady=3)
        self.search_entry.bind("<Return>", lambda e: self._do_search())
        
        ModernButton(self.search_frame, icon="◀", width=32, command=lambda: self._nav_search(-1)).pack(side=tk.LEFT, padx=2)
        ModernButton(self.search_frame, icon="▶", width=32, command=lambda: self._nav_search(1)).pack(side=tk.LEFT, padx=2)
        
        self.search_results_label = tk.Label(self.search_frame, text="", bg=Theme.BG_TERTIARY, fg=Theme.FG_SECONDARY,
                                             font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.search_results_label.pack(side=tk.LEFT, padx=Theme.PAD_MD)
        
        ModernButton(self.search_frame, icon="✕", width=32, command=self._hide_search).pack(side=tk.RIGHT, padx=Theme.PAD_MD)
        
        # Canvas
        canvas_frame = tk.Frame(canvas_container, bg=Theme.BG_CANVAS)
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        
        self.h_scroll = ttk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        self.v_scroll = ttk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
        
        self.canvas = tk.Canvas(canvas_frame, bg=Theme.BG_CANVAS, highlightthickness=0,
                               xscrollcommand=self.h_scroll.set, yscrollcommand=self.v_scroll.set)
        
        self.h_scroll.configure(command=self.canvas.xview)
        self.v_scroll.configure(command=self.canvas.yview)
        
        self.v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        
        # Canvas bindings
        self.canvas.bind("<Button-1>", self._canvas_click)
        self.canvas.bind("<B1-Motion>", self._canvas_drag)
        self.canvas.bind("<ButtonRelease-1>", self._canvas_release)
        self.canvas.bind("<MouseWheel>", self._canvas_scroll)
        self.canvas.bind("<Button-3>", self._canvas_context)
    
    def _build_properties_panel(self, parent):
        self.props_panel = tk.Frame(parent, bg=Theme.BG_SECONDARY, width=220)
        self.props_panel.pack(side=tk.RIGHT, fill=tk.Y)
        self.props_panel.pack_propagate(False)
        
        # Header
        header = tk.Frame(self.props_panel, bg=Theme.BG_TERTIARY)
        header.pack(fill=tk.X)
        tk.Label(header, text="Properties", bg=Theme.BG_TERTIARY, fg=Theme.FG_PRIMARY,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_MD, "bold"),
                padx=Theme.PAD_MD, pady=Theme.PAD_MD).pack(anchor="w")
        
        self.props_content = tk.Frame(self.props_panel, bg=Theme.BG_SECONDARY)
        self.props_content.pack(fill=tk.BOTH, expand=True, padx=Theme.PAD_MD, pady=Theme.PAD_MD)
    
    def _build_status_bar(self, parent):
        status = tk.Frame(parent, bg=Theme.BG_PRIMARY, height=Theme.STATUSBAR_HEIGHT)
        status.pack(fill=tk.X, side=tk.BOTTOM)
        status.pack_propagate(False)
        
        self.status_left = tk.Label(status, text="Ready", bg=Theme.BG_PRIMARY, fg=Theme.FG_SECONDARY,
                                    font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.status_left.pack(side=tk.LEFT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
        
        self.status_right = tk.Label(status, text="", bg=Theme.BG_PRIMARY, fg=Theme.FG_SECONDARY,
                                     font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        self.status_right.pack(side=tk.RIGHT, padx=Theme.PAD_MD, pady=Theme.PAD_SM)
    
    def _bind_shortcuts(self):
        shortcuts = [
            ("<Control-n>", self._new_doc), ("<Control-o>", self._open_doc),
            ("<Control-s>", self._save_doc), ("<Control-w>", self._close_tab),
            ("<Control-f>", self._show_search),
            ("<Control-plus>", self._zoom_in), ("<Control-minus>", self._zoom_out),
            ("<Control-equal>", self._zoom_in), ("<Control-0>", self._zoom_fit),
            ("<Home>", self._first_page), ("<End>", self._last_page),
            ("<Prior>", self._prev_page), ("<Next>", self._next_page),
            ("<Delete>", self._delete_page), ("<Escape>", lambda: self._set_tool(ToolMode.SELECT)),
        ]
        for key, cmd in shortcuts:
            self.bind(key, lambda e, c=cmd: c())
    
    # =========================================================================
    # SIDEBAR
    # =========================================================================
    
    def _show_sidebar_content(self, key):
        for k, tab in self.sidebar_tabs.items():
            tab.set_active(k == key)
        
        for panel in [self.pages_panel, self.bookmarks_panel, self.comments_panel]:
            panel.pack_forget()
        
        if key == "pages":
            self.pages_panel.pack(fill=tk.BOTH, expand=True)
        elif key == "bookmarks":
            self.bookmarks_panel.pack(fill=tk.BOTH, expand=True)
            self._refresh_bookmarks()
        elif key == "comments":
            self.comments_panel.pack(fill=tk.BOTH, expand=True)
            self._refresh_comments()
        
        self.sidebar_mode = key
    
    def _refresh_thumbnails(self):
        for t in self.thumbnails:
            t.destroy()
        self.thumbnails = []
        
        if not self.doc:
            return
        
        for i in range(self.doc.page_count):
            self._create_thumbnail(i)
    
    def _create_thumbnail(self, page_num):
        img = self.doc.render_page(page_num, zoom=0.15)
        if not img:
            return
        
        img.thumbnail((120, 160), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(img)
        
        frame = tk.Frame(self.thumb_frame, bg=Theme.BG_SECONDARY, cursor="hand2")
        frame.pack(pady=Theme.PAD_SM, padx=Theme.PAD_SM)
        
        canvas = tk.Canvas(frame, width=130, height=170, bg=Theme.BG_SECONDARY, highlightthickness=0)
        canvas.pack()
        
        # Thumbnail with border
        border_color = Theme.ACCENT if page_num == self.current_page else Theme.BORDER_LIGHT
        canvas.create_rectangle(9, 9, 121, 151, fill="white", outline=border_color, width=2)
        canvas.create_image(65, 80, image=photo)
        canvas.create_text(65, 162, text=str(page_num + 1), fill=Theme.FG_SECONDARY,
                          font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        
        canvas.image = photo
        canvas.bind("<Button-1>", lambda e, p=page_num: self._goto_page(p))
        canvas.bind("<Button-3>", lambda e, p=page_num: self._page_context(e, p))
        
        self.thumbnails.append(frame)
    
    def _update_thumbnail_selection(self):
        for i, thumb in enumerate(self.thumbnails):
            canvas = thumb.winfo_children()[0]
            border_color = Theme.ACCENT if i == self.current_page else Theme.BORDER_LIGHT
            canvas.itemconfigure(1, outline=border_color)
    
    def _refresh_bookmarks(self):
        self.bookmarks_tree.delete(*self.bookmarks_tree.get_children())
        if not self.doc:
            return
        
        bookmarks = self.doc.get_bookmarks()
        parents = {0: ""}
        for level, title, page in bookmarks:
            parent = parents.get(level - 1, "")
            item = self.bookmarks_tree.insert(parent, "end", text=f"{title} ({page + 1})", values=(page,))
            parents[level] = item
    
    def _refresh_comments(self):
        self.comments_list.delete(0, tk.END)
        if not self.doc:
            return
        for c in self.doc.comments:
            preview = c.content[:35] + "..." if len(c.content) > 35 else c.content
            self.comments_list.insert(tk.END, f"p.{c.page + 1}: {preview}")
    
    def _on_bookmark_select(self, e):
        sel = self.bookmarks_tree.selection()
        if sel:
            item = self.bookmarks_tree.item(sel[0])
            if item['values']:
                self._goto_page(item['values'][0])
    
    def _on_comment_select(self, e):
        sel = self.comments_list.curselection()
        if sel and self.doc and sel[0] < len(self.doc.comments):
            self._goto_page(self.doc.comments[sel[0]].page)
    
    # =========================================================================
    # DOCUMENT MANAGEMENT
    # =========================================================================
    
    def _new_doc(self):
        doc_id = f"doc_{len(self.documents)}_{datetime.now().timestamp()}"
        doc = PDFDocument()
        doc.create_new()
        self.documents[doc_id] = doc
        self._add_tab(doc_id, "Untitled")
        self._switch_to_doc(doc_id)
        self._status("New document created")
    
    def _open_doc(self, filepath=None):
        if not filepath:
            filepath = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")])
        if not filepath:
            return
        
        doc_id = f"doc_{len(self.documents)}_{datetime.now().timestamp()}"
        doc = PDFDocument()
        
        if doc.open(filepath):
            self.documents[doc_id] = doc
            self._add_tab(doc_id, doc.filename)
            self._switch_to_doc(doc_id)
            self._add_recent(filepath)
            self._status(f"Opened: {doc.filename}")
        else:
            messagebox.showerror("Error", "Failed to open PDF")
    
    def _save_doc(self):
        if not self.doc:
            return
        if not self.doc.filepath:
            self._save_as()
            return
        if self.doc.save():
            self._status("Saved")
            self._update_tab_title()
        else:
            messagebox.showerror("Error", "Failed to save")
    
    def _save_as(self):
        if not self.doc:
            return
        filepath = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if filepath:
            if self.doc.save(filepath):
                self._status(f"Saved: {self.doc.filename}")
                self._update_tab_title()
                self._add_recent(filepath)
    
    def _close_tab(self):
        if not self.active_doc_id:
            return
        doc = self.doc
        if doc and doc.is_modified:
            r = messagebox.askyesnocancel("Save Changes?", f"Save changes to {doc.filename}?")
            if r is None:
                return
            if r:
                self._save_doc()
        
        self._remove_tab(self.active_doc_id)
        if self.active_doc_id in self.documents:
            self.documents[self.active_doc_id].close()
            del self.documents[self.active_doc_id]
        
        if self.tabs:
            self._switch_to_doc(list(self.tabs.keys())[0])
        else:
            self.active_doc_id = None
            self.current_page = 0
            self._show_welcome()
    
    def _add_tab(self, doc_id, title):
        tab = TabButton(self.tab_bar, title=title, doc_id=doc_id,
                       on_select=self._switch_to_doc, on_close=self._close_tab_by_id)
        tab.pack(side=tk.LEFT, padx=1)
        self.tabs[doc_id] = tab
    
    def _remove_tab(self, doc_id):
        if doc_id in self.tabs:
            self.tabs[doc_id].destroy()
            del self.tabs[doc_id]
    
    def _close_tab_by_id(self, doc_id):
        old_active = self.active_doc_id
        self.active_doc_id = doc_id
        self._close_tab()
        if self.tabs and old_active != doc_id:
            self._switch_to_doc(old_active)
    
    def _switch_to_doc(self, doc_id):
        if doc_id not in self.documents:
            return
        self.active_doc_id = doc_id
        self.current_page = 0
        self.zoom = 1.0
        
        for did, tab in self.tabs.items():
            tab.set_active(did == doc_id)
        
        self._refresh_all()
    
    def _update_tab_title(self):
        if self.active_doc_id in self.tabs and self.doc:
            title = self.doc.filename + (" *" if self.doc.is_modified else "")
            self.tabs[self.active_doc_id].set_title(title)
            self.title(f"PDF Editor Pro - {title}")
    
    def _add_recent(self, filepath):
        recent = self.config_data.get("recent_files", [])
        if filepath in recent:
            recent.remove(filepath)
        recent.insert(0, filepath)
        self.config_data["recent_files"] = recent[:Config.MAX_RECENT_FILES]
        Config.save(self.config_data)
    
    # =========================================================================
    # VIEW & RENDERING
    # =========================================================================
    
    def _refresh_all(self):
        self._render_page()
        self._refresh_thumbnails()
        self._refresh_bookmarks()
        self._refresh_comments()
        self._update_properties()
        self._update_ui()
    
    def _render_page(self):
        if not self.doc:
            self._show_welcome()
            return
        
        img = self.doc.render_page(self.current_page, self.zoom)
        if not img:
            return
        
        self.page_image = ImageTk.PhotoImage(img)
        self.canvas.delete("all")
        
        cw = self.canvas.winfo_width() or 800
        ch = self.canvas.winfo_height() or 600
        iw, ih = img.size
        
        x = max(cw // 2, iw // 2)
        y = max(ch // 2, ih // 2)
        
        # Shadow
        self.canvas.create_rectangle(x - iw//2 + 6, y - ih//2 + 6,
                                     x + iw//2 + 6, y + ih//2 + 6,
                                     fill=Theme.SHADOW, outline="")
        
        # Page background
        self.canvas.create_rectangle(x - iw//2, y - ih//2, x + iw//2, y + ih//2,
                                     fill="white", outline=Theme.BORDER_DARK)
        
        # Page image
        self.canvas.create_image(x, y, image=self.page_image)
        
        self.img_offset = (x - iw // 2, y - ih // 2)
        
        # Draw comments
        for c in self.doc.comments:
            if c.page == self.current_page:
                cx = self.img_offset[0] + c.x * self.zoom
                cy = self.img_offset[1] + c.y * self.zoom
                self.canvas.create_polygon(cx, cy, cx+18, cy, cx+18, cy+22,
                                          cx+9, cy+15, cx, cy+15,
                                          fill=c.color, outline=Theme.BORDER_DARK)
        
        # Search highlights
        for sr in self.search_results:
            if sr.page == self.current_page:
                r = sr.rect
                x1 = self.img_offset[0] + r[0] * self.zoom
                y1 = self.img_offset[1] + r[1] * self.zoom
                x2 = self.img_offset[0] + r[2] * self.zoom
                y2 = self.img_offset[1] + r[3] * self.zoom
                self.canvas.create_rectangle(x1, y1, x2, y2, fill=Theme.HIGHLIGHT,
                                            stipple="gray50", outline="")
        
        self.canvas.configure(scrollregion=(0, 0, max(cw, iw+100), max(ch, ih+100)))
    
    def _show_welcome(self):
        self.canvas.delete("all")
        cx, cy = 500, 350
        
        self.canvas.create_text(cx, cy - 80, text="📄", font=(Theme.FONT_FAMILY, 64), fill=Theme.ACCENT)
        self.canvas.create_text(cx, cy, text="PDF Editor Pro",
                               font=(Theme.FONT_FAMILY, 32, "bold"), fill=Theme.FG_PRIMARY)
        self.canvas.create_text(cx, cy + 45, text="Professional PDF Editing Suite",
                               font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_LG), fill=Theme.FG_SECONDARY)
        
        recent = self.config_data.get("recent_files", [])[:5]
        if recent:
            self.canvas.create_text(cx, cy + 110, text="Recent Files",
                                   font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_MD, "bold"), fill=Theme.FG_PRIMARY)
            for i, path in enumerate(recent):
                y = cy + 140 + i * 26
                name = os.path.basename(path)
                tag = f"recent_{i}"
                self.canvas.create_text(cx, y, text=name, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM),
                                       fill=Theme.ACCENT_LIGHT, tags=tag)
                self.canvas.tag_bind(tag, "<Button-1>", lambda e, p=path: self._open_doc(p))
                self.canvas.tag_bind(tag, "<Enter>", lambda e, t=tag: self.canvas.itemconfigure(t, fill=Theme.FG_PRIMARY))
                self.canvas.tag_bind(tag, "<Leave>", lambda e, t=tag: self.canvas.itemconfigure(t, fill=Theme.ACCENT_LIGHT))
    
    def _update_ui(self):
        self.page_entry.delete(0, tk.END)
        self.page_entry.insert(0, str(self.current_page + 1) if self.doc else "0")
        self.page_total.configure(text=f"/ {self.doc.page_count if self.doc else 0}")
        self.zoom_label.configure(text=f"{int(self.zoom * 100)}%")
        
        if self.doc:
            mod = " *" if self.doc.is_modified else ""
            self.status_right.configure(text=f"Page {self.current_page + 1} of {self.doc.page_count}{mod}")
            self._update_tab_title()
        else:
            self.title("PDF Editor Pro")
    
    def _update_properties(self):
        for w in self.props_content.winfo_children():
            w.destroy()
        
        if not self.doc:
            return
        
        page = self.doc.get_page(self.current_page)
        if not page:
            return
        
        # Page info
        tk.Label(self.props_content, text="Page", bg=Theme.BG_SECONDARY, fg=Theme.ACCENT_LIGHT,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM, "bold")).pack(anchor="w", pady=(0, Theme.PAD_SM))
        
        info = [
            ("Number", str(self.current_page + 1)),
            ("Width", f"{page.rect.width:.0f} pt"),
            ("Height", f"{page.rect.height:.0f} pt"),
            ("Rotation", f"{page.rotation}°"),
        ]
        
        for label, value in info:
            row = tk.Frame(self.props_content, bg=Theme.BG_SECONDARY)
            row.pack(fill=tk.X, pady=2)
            tk.Label(row, text=label, bg=Theme.BG_SECONDARY, fg=Theme.FG_MUTED,
                    font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM), width=10, anchor="w").pack(side=tk.LEFT)
            tk.Label(row, text=value, bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY,
                    font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM)).pack(side=tk.LEFT)
    
    def _status(self, msg):
        self.status_left.configure(text=msg)
    
    # =========================================================================
    # NAVIGATION
    # =========================================================================
    
    def _first_page(self):
        if self.doc:
            self._goto_page(0)
    
    def _prev_page(self):
        if self.doc and self.current_page > 0:
            self._goto_page(self.current_page - 1)
    
    def _next_page(self):
        if self.doc and self.current_page < self.doc.page_count - 1:
            self._goto_page(self.current_page + 1)
    
    def _last_page(self):
        if self.doc:
            self._goto_page(self.doc.page_count - 1)
    
    def _goto_page(self, page_num):
        if self.doc and 0 <= page_num < self.doc.page_count:
            self.current_page = page_num
            self._render_page()
            self._update_thumbnail_selection()
            self._update_properties()
            self._update_ui()
    
    def _goto_page_entry(self, e=None):
        try:
            p = int(self.page_entry.get()) - 1
            self._goto_page(p)
        except:
            pass
    
    # =========================================================================
    # ZOOM
    # =========================================================================
    
    def _zoom_in(self):
        self.zoom = min(Config.MAX_ZOOM, self.zoom * 1.25)
        self._render_page()
        self._update_ui()
    
    def _zoom_out(self):
        self.zoom = max(Config.MIN_ZOOM, self.zoom / 1.25)
        self._render_page()
        self._update_ui()
    
    def _zoom_100(self):
        self.zoom = 1.0
        self._render_page()
        self._update_ui()
    
    def _zoom_fit(self):
        if not self.doc:
            return
        pw, ph = self.doc.get_page_size(self.current_page)
        cw = self.canvas.winfo_width() - 60
        ch = self.canvas.winfo_height() - 60
        self.zoom = min(cw / pw, ch / ph, Config.MAX_ZOOM)
        self._render_page()
        self._update_ui()
    
    def _canvas_scroll(self, e):
        if e.state & 0x4:  # Ctrl
            self._zoom_in() if e.delta > 0 else self._zoom_out()
        else:
            self.canvas.yview_scroll(-1 * (e.delta // 120), "units")
    
    # =========================================================================
    # TOOLS
    # =========================================================================
    
    def _set_tool(self, mode):
        self.tool_mode = mode
        for m, btn in self.tool_buttons.items():
            btn.set_active(m == mode)
        
        cursors = {
            ToolMode.SELECT: "arrow", ToolMode.PAN: "fleur", ToolMode.TEXT: "xterm",
            ToolMode.STICKY_NOTE: "plus", ToolMode.DRAW: "pencil", ToolMode.CROP: "cross",
        }
        self.canvas.configure(cursor=cursors.get(mode, "cross"))
        self._status(f"Tool: {mode.name.replace('_', ' ').title()}")
    
    def _canvas_to_pdf(self, cx, cy):
        if not hasattr(self, 'img_offset'):
            return 0, 0
        return (cx - self.img_offset[0]) / self.zoom, (cy - self.img_offset[1]) / self.zoom
    
    def _canvas_click(self, e):
        if not self.doc:
            return
        
        cx = self.canvas.canvasx(e.x)
        cy = self.canvas.canvasy(e.y)
        self.drag_start = (cx, cy)
        self.draw_points = [(cx, cy)]
        
        px, py = self._canvas_to_pdf(cx, cy)
        
        if self.tool_mode == ToolMode.TEXT:
            self._text_dialog(px, py)
        elif self.tool_mode == ToolMode.STICKY_NOTE:
            self._comment_dialog(px, py)
        elif self.tool_mode == ToolMode.STAMP and self.selected_stamp:
            self.doc.add_stamp(self.current_page, px, py, self.selected_stamp)
            self._render_page()
    
    def _canvas_drag(self, e):
        if not self.doc or not self.drag_start:
            return
        
        cx = self.canvas.canvasx(e.x)
        cy = self.canvas.canvasy(e.y)
        
        if self.tool_mode == ToolMode.PAN:
            dx = cx - self.drag_start[0]
            dy = cy - self.drag_start[1]
            self.canvas.xview_scroll(int(-dx/15), "units")
            self.canvas.yview_scroll(int(-dy/15), "units")
        elif self.tool_mode == ToolMode.DRAW:
            self.draw_points.append((cx, cy))
            if len(self.draw_points) >= 2:
                self.canvas.create_line(self.draw_points[-2][0], self.draw_points[-2][1],
                                       cx, cy, fill="#000000", width=2, tags="temp")
        elif self.tool_mode in (ToolMode.RECTANGLE, ToolMode.CIRCLE, ToolMode.LINE,
                               ToolMode.ARROW, ToolMode.HIGHLIGHT, ToolMode.UNDERLINE,
                               ToolMode.STRIKETHROUGH, ToolMode.REDACT, ToolMode.CROP):
            self.canvas.delete("temp")
            x1, y1 = self.drag_start
            
            if self.tool_mode == ToolMode.RECTANGLE:
                self.canvas.create_rectangle(x1, y1, cx, cy, outline="#000000", width=2, tags="temp")
            elif self.tool_mode == ToolMode.CIRCLE:
                self.canvas.create_oval(x1, y1, cx, cy, outline="#000000", width=2, tags="temp")
            elif self.tool_mode == ToolMode.LINE:
                self.canvas.create_line(x1, y1, cx, cy, fill="#000000", width=2, tags="temp")
            elif self.tool_mode == ToolMode.ARROW:
                self.canvas.create_line(x1, y1, cx, cy, fill="#000000", width=2, arrow=tk.LAST, tags="temp")
            elif self.tool_mode in (ToolMode.HIGHLIGHT, ToolMode.UNDERLINE, ToolMode.STRIKETHROUGH):
                self.canvas.create_rectangle(x1, y1, cx, cy, fill=Theme.HIGHLIGHT, stipple="gray50", outline="", tags="temp")
            elif self.tool_mode == ToolMode.REDACT:
                self.canvas.create_rectangle(x1, y1, cx, cy, fill="black", outline="", tags="temp")
            elif self.tool_mode == ToolMode.CROP:
                self.canvas.create_rectangle(x1, y1, cx, cy, outline=Theme.ACCENT, width=2, dash=(4, 4), tags="temp")
    
    def _canvas_release(self, e):
        if not self.doc or not self.drag_start:
            return
        
        cx = self.canvas.canvasx(e.x)
        cy = self.canvas.canvasy(e.y)
        
        x1, y1 = self._canvas_to_pdf(*self.drag_start)
        x2, y2 = self._canvas_to_pdf(cx, cy)
        rect = (min(x1, x2), min(y1, y2), max(x1, x2), max(y1, y2))
        
        self.canvas.delete("temp")
        
        if self.tool_mode == ToolMode.DRAW and len(self.draw_points) >= 2:
            pts = [self._canvas_to_pdf(p[0], p[1]) for p in self.draw_points]
            self.doc.add_freehand(self.current_page, pts)
            self._render_page()
        elif self.tool_mode == ToolMode.RECTANGLE:
            self.doc.add_rect(self.current_page, rect)
            self._render_page()
        elif self.tool_mode == ToolMode.CIRCLE:
            self.doc.add_circle(self.current_page, rect)
            self._render_page()
        elif self.tool_mode == ToolMode.LINE:
            self.doc.add_line(self.current_page, (x1, y1), (x2, y2))
            self._render_page()
        elif self.tool_mode == ToolMode.ARROW:
            self.doc.add_arrow(self.current_page, (x1, y1), (x2, y2))
            self._render_page()
        elif self.tool_mode == ToolMode.HIGHLIGHT:
            self.doc.add_highlight(self.current_page, rect)
            self._render_page()
        elif self.tool_mode == ToolMode.UNDERLINE:
            self.doc.add_underline(self.current_page, rect)
            self._render_page()
        elif self.tool_mode == ToolMode.STRIKETHROUGH:
            self.doc.add_strikethrough(self.current_page, rect)
            self._render_page()
        elif self.tool_mode == ToolMode.REDACT:
            if messagebox.askyesno("Redact", "Permanently redact this area?"):
                self.doc.redact_area(self.current_page, rect)
                self._render_page()
        elif self.tool_mode == ToolMode.CROP:
            if messagebox.askyesno("Crop", "Crop page to selected area?"):
                self.doc.crop_page(self.current_page, rect)
                self._refresh_all()
        
        self.drag_start = None
        self.draw_points = []
        self._update_ui()
    
    def _canvas_context(self, e):
        if not self.doc:
            return
        px, py = self._canvas_to_pdf(self.canvas.canvasx(e.x), self.canvas.canvasy(e.y))
        
        menu = tk.Menu(self, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY)
        menu.add_command(label="Add Text", command=lambda: self._text_dialog(px, py))
        menu.add_command(label="Add Comment", command=lambda: self._comment_dialog(px, py))
        menu.add_separator()
        menu.add_command(label="Copy Page Text", command=self._copy_text)
        menu.tk_popup(e.x_root, e.y_root)
    
    def _page_context(self, e, page_num):
        menu = tk.Menu(self, tearoff=0, bg=Theme.BG_ELEVATED, fg=Theme.FG_PRIMARY)
        menu.add_command(label="Insert Page Before", command=lambda: self._insert_page_at(page_num))
        menu.add_command(label="Insert Page After", command=lambda: self._insert_page_at(page_num + 1))
        menu.add_command(label="Duplicate", command=lambda: self._duplicate_page_at(page_num))
        menu.add_separator()
        menu.add_command(label="Rotate CW", command=lambda: self._rotate_page(page_num, 90))
        menu.add_command(label="Rotate CCW", command=lambda: self._rotate_page(page_num, -90))
        menu.add_separator()
        menu.add_command(label="Delete", command=lambda: self._delete_page_at(page_num))
        menu.tk_popup(e.x_root, e.y_root)
    
    # =========================================================================
    # DIALOGS
    # =========================================================================
    
    def _create_dialog(self, title, width=400, height=300):
        dialog = tk.Toplevel(self)
        dialog.title(title)
        dialog.geometry(f"{width}x{height}")
        dialog.configure(bg=Theme.BG_SECONDARY)
        dialog.transient(self)
        dialog.grab_set()
        
        # Center on parent
        dialog.update_idletasks()
        x = self.winfo_x() + (self.winfo_width() - width) // 2
        y = self.winfo_y() + (self.winfo_height() - height) // 2
        dialog.geometry(f"+{x}+{y}")
        
        return dialog
    
    def _text_dialog(self, x, y):
        dialog = self._create_dialog("Add Text", 420, 240)
        
        tk.Label(dialog, text="Enter text:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM)).pack(pady=(Theme.PAD_LG, Theme.PAD_SM))
        
        text_box = tk.Text(dialog, height=4, width=45, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY,
                          insertbackground=Theme.FG_PRIMARY, relief=tk.FLAT,
                          font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM))
        text_box.pack(padx=Theme.PAD_LG, pady=Theme.PAD_SM)
        text_box.focus_set()
        
        opt = tk.Frame(dialog, bg=Theme.BG_SECONDARY)
        opt.pack(pady=Theme.PAD_SM)
        tk.Label(opt, text="Size:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(side=tk.LEFT)
        size_var = tk.StringVar(value="12")
        tk.Spinbox(opt, from_=6, to=72, textvariable=size_var, width=5, bg=Theme.BG_INPUT,
                  fg=Theme.FG_PRIMARY).pack(side=tk.LEFT, padx=Theme.PAD_SM)
        
        def add():
            text = text_box.get("1.0", tk.END).strip()
            if text:
                self.doc.add_text(self.current_page, text, x, y, int(size_var.get()))
                self._render_page()
            dialog.destroy()
        
        btn_frame = tk.Frame(dialog, bg=Theme.BG_SECONDARY)
        btn_frame.pack(pady=Theme.PAD_LG)
        ModernButton(btn_frame, text="Add", command=add, style="primary", width=100).pack(side=tk.LEFT, padx=Theme.PAD_SM)
        ModernButton(btn_frame, text="Cancel", command=dialog.destroy, width=100).pack(side=tk.LEFT)
    
    def _comment_dialog(self, x, y):
        dialog = self._create_dialog("Add Comment", 380, 200)
        
        tk.Label(dialog, text="Comment:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_SM)).pack(pady=(Theme.PAD_LG, Theme.PAD_SM))
        
        text_box = tk.Text(dialog, height=4, width=38, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY,
                          insertbackground=Theme.FG_PRIMARY, relief=tk.FLAT)
        text_box.pack(padx=Theme.PAD_LG, pady=Theme.PAD_SM)
        text_box.focus_set()
        
        def add():
            text = text_box.get("1.0", tk.END).strip()
            if text:
                self.doc.add_comment(self.current_page, x, y, text)
                self._render_page()
                self._refresh_comments()
            dialog.destroy()
        
        btn_frame = tk.Frame(dialog, bg=Theme.BG_SECONDARY)
        btn_frame.pack(pady=Theme.PAD_LG)
        ModernButton(btn_frame, text="Add", command=add, style="primary", width=100).pack(side=tk.LEFT, padx=Theme.PAD_SM)
        ModernButton(btn_frame, text="Cancel", command=dialog.destroy, width=100).pack(side=tk.LEFT)
    
    def _show_stamp_dialog(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Select Stamp", 420, 340)
        
        tk.Label(dialog, text="Select a Stamp", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY,
                font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_LG, "bold")).pack(pady=Theme.PAD_LG)
        
        frame = tk.Frame(dialog, bg=Theme.BG_SECONDARY)
        frame.pack(fill=tk.BOTH, expand=True, padx=Theme.PAD_LG, pady=Theme.PAD_SM)
        
        def select(stamp):
            self.selected_stamp = stamp
            self._set_tool(ToolMode.STAMP)
            dialog.destroy()
            self._status(f"Stamp: {stamp['name']} - Click to place")
        
        for i, stamp in enumerate(BUILTIN_STAMPS):
            btn = tk.Button(frame, text=stamp['text'], bg=stamp['bg'], fg=stamp['fg'],
                           font=(Theme.FONT_FAMILY, 10, "bold"), relief=tk.FLAT, padx=10, pady=6,
                           command=lambda s=stamp: select(s))
            btn.grid(row=i//3, column=i%3, padx=4, pady=4, sticky='ew')
        
        for i in range(3):
            frame.columnconfigure(i, weight=1)
        
        ModernButton(dialog, text="Cancel", command=dialog.destroy, width=100).pack(pady=Theme.PAD_LG)
    
    def _watermark_dialog(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Add Watermark", 380, 280)
        
        tk.Label(dialog, text="Watermark Text:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_LG, Theme.PAD_SM))
        text_entry = ModernEntry(dialog, width=30)
        text_entry.pack(ipady=4)
        text_entry.insert(0, "CONFIDENTIAL")
        
        tk.Label(dialog, text="Font Size:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_SM))
        size_var = tk.StringVar(value="48")
        tk.Spinbox(dialog, from_=12, to=144, textvariable=size_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        tk.Label(dialog, text="Rotation:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_SM))
        angle_var = tk.StringVar(value="45")
        tk.Spinbox(dialog, from_=-90, to=90, textvariable=angle_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        def apply():
            self.doc.add_watermark(text_entry.get(), int(size_var.get()), angle=float(angle_var.get()))
            self._render_page()
            dialog.destroy()
            self._status("Watermark added")
        
        ModernButton(dialog, text="Apply to All Pages", command=apply, style="primary", width=160).pack(pady=Theme.PAD_LG)
    
    def _header_footer_dialog(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Headers & Footers", 450, 320)
        
        tk.Label(dialog, text="Placeholders: {page}, {pages}, {date}", bg=Theme.BG_SECONDARY,
                fg=Theme.FG_MUTED, font=(Theme.FONT_FAMILY, Theme.FONT_SIZE_XS)).pack(pady=(Theme.PAD_MD, Theme.PAD_SM))
        
        tk.Label(dialog, text="Header:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_SM, Theme.PAD_XS))
        header_entry = ModernEntry(dialog, width=45)
        header_entry.pack(ipady=4)
        
        tk.Label(dialog, text="Footer:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        footer_entry = ModernEntry(dialog, width=45)
        footer_entry.pack(ipady=4)
        footer_entry.insert(0, "Page {page} of {pages}")
        
        tk.Label(dialog, text="Font Size:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        size_var = tk.StringVar(value="10")
        tk.Spinbox(dialog, from_=6, to=24, textvariable=size_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        def apply():
            h = header_entry.get() or None
            f = footer_entry.get() or None
            if h or f:
                self.doc.add_header_footer(h, f, int(size_var.get()))
                self._render_page()
                self._status("Header/footer added")
            dialog.destroy()
        
        ModernButton(dialog, text="Apply", command=apply, style="primary", width=120).pack(pady=Theme.PAD_LG)
    
    def _bates_dialog(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Bates Numbering", 380, 340)
        
        tk.Label(dialog, text="Prefix:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_LG, Theme.PAD_XS))
        prefix_entry = ModernEntry(dialog, width=20)
        prefix_entry.pack(ipady=3)
        prefix_entry.insert(0, "DOC-")
        
        tk.Label(dialog, text="Start Number:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        start_var = tk.StringVar(value="1")
        tk.Spinbox(dialog, from_=1, to=999999, textvariable=start_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        tk.Label(dialog, text="Digits:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        digits_var = tk.StringVar(value="6")
        tk.Spinbox(dialog, from_=3, to=10, textvariable=digits_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        tk.Label(dialog, text="Position:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        pos_var = tk.StringVar(value="bottom-right")
        ttk.Combobox(dialog, textvariable=pos_var, values=["top-left", "top-right", "bottom-left", "bottom-right"], width=15).pack()
        
        def apply():
            self.doc.add_bates_numbers(prefix_entry.get(), int(start_var.get()), int(digits_var.get()), pos_var.get())
            self._render_page()
            dialog.destroy()
            self._status("Bates numbers added")
        
        ModernButton(dialog, text="Apply", command=apply, style="primary", width=120).pack(pady=Theme.PAD_LG)
    
    def _password_dialog(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Password Protection", 350, 200)
        
        tk.Label(dialog, text="Password:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_LG, Theme.PAD_XS))
        pass_entry = tk.Entry(dialog, show="*", width=25, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY, relief=tk.FLAT)
        pass_entry.pack(ipady=4)
        
        tk.Label(dialog, text="Confirm:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        confirm_entry = tk.Entry(dialog, show="*", width=25, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY, relief=tk.FLAT)
        confirm_entry.pack(ipady=4)
        
        def apply():
            if pass_entry.get() != confirm_entry.get():
                messagebox.showerror("Error", "Passwords don't match")
                return
            output = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
            if output:
                try:
                    self.doc.doc.save(output, encryption=fitz.PDF_ENCRYPT_AES_256, user_pw=pass_entry.get())
                    dialog.destroy()
                    self._status("Protected PDF saved")
                except Exception as e:
                    messagebox.showerror("Error", str(e))
        
        ModernButton(dialog, text="Save Protected", command=apply, style="primary", width=140).pack(pady=Theme.PAD_LG)
    
    # =========================================================================
    # SEARCH
    # =========================================================================
    
    def _show_search(self):
        self.search_frame.pack(fill=tk.X, before=self.canvas.master.master)
        self.search_entry.focus_set()
    
    def _hide_search(self):
        self.search_frame.pack_forget()
        self.search_results = []
        self._render_page()
    
    def _do_search(self):
        query = self.search_entry.get_value()
        if not query or not self.doc:
            return
        
        self.search_results = self.doc.search_text(query)
        self.search_idx = 0
        
        if self.search_results:
            self.search_results_label.configure(text=f"1 of {len(self.search_results)}")
            self._goto_page(self.search_results[0].page)
        else:
            self.search_results_label.configure(text="No results")
            self._render_page()
    
    def _nav_search(self, direction):
        if not self.search_results:
            return
        self.search_idx = (self.search_idx + direction) % len(self.search_results)
        self.search_results_label.configure(text=f"{self.search_idx + 1} of {len(self.search_results)}")
        self._goto_page(self.search_results[self.search_idx].page)
    
    # =========================================================================
    # PAGE OPERATIONS
    # =========================================================================
    
    def _insert_page(self):
        if self.doc:
            self.doc.insert_page(self.current_page + 1)
            self._refresh_all()
    
    def _insert_page_at(self, index):
        if self.doc:
            self.doc.insert_page(index)
            self._refresh_all()
    
    def _duplicate_page(self):
        self._duplicate_page_at(self.current_page)
    
    def _duplicate_page_at(self, page_num):
        if self.doc:
            self.doc.duplicate_page(page_num)
            self._refresh_all()
    
    def _delete_page(self):
        self._delete_page_at(self.current_page)
    
    def _delete_page_at(self, page_num):
        if not self.doc or self.doc.page_count <= 1:
            messagebox.showwarning("Warning", "Cannot delete the only page")
            return
        if messagebox.askyesno("Delete Page", f"Delete page {page_num + 1}?"):
            self.doc.delete_page(page_num)
            if self.current_page >= self.doc.page_count:
                self.current_page = self.doc.page_count - 1
            self._refresh_all()
    
    def _extract_page(self):
        if not self.doc:
            return
        output = filedialog.asksaveasfilename(defaultextension=".pdf", initialname=f"page_{self.current_page+1}.pdf")
        if output:
            new_doc = fitz.open()
            new_doc.insert_pdf(self.doc.doc, from_page=self.current_page, to_page=self.current_page)
            new_doc.save(output)
            new_doc.close()
            self._status(f"Page extracted to {os.path.basename(output)}")
    
    def _rotate(self, angle):
        self._rotate_page(self.current_page, angle)
    
    def _rotate_page(self, page_num, angle):
        if self.doc:
            self.doc.rotate_page(page_num, angle)
            self._refresh_all()
    
    # =========================================================================
    # DOCUMENT OPERATIONS
    # =========================================================================
    
    def _add_image(self):
        if not self.doc:
            return
        filepath = filedialog.askopenfilename(filetypes=[("Images", "*.png *.jpg *.jpeg *.gif *.bmp")])
        if filepath:
            if self.doc.add_image(self.current_page, filepath):
                self._render_page()
                self._status("Image added")
    
    def _copy_text(self):
        if self.doc:
            text = self.doc.get_text(self.current_page)
            self.clipboard_clear()
            self.clipboard_append(text)
            self._status("Text copied to clipboard")
    
    def _merge_pdfs(self):
        files = filedialog.askopenfilenames(filetypes=[("PDF", "*.pdf")])
        if not files:
            return
        output = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not output:
            return
        
        merged = fitz.open()
        for f in files:
            doc = fitz.open(f)
            merged.insert_pdf(doc)
            doc.close()
        merged.save(output)
        merged.close()
        
        if messagebox.askyesno("Done", f"Merged {len(files)} files. Open result?"):
            self._open_doc(output)
    
    def _split_doc(self):
        if not self.doc:
            return
        output_dir = filedialog.askdirectory(title="Select output folder")
        if output_dir:
            files = self.doc.split_pages(output_dir)
            messagebox.showinfo("Done", f"Split into {len(files)} files")
    
    def _compress_doc(self):
        if not self.doc:
            return
        output = filedialog.asksaveasfilename(defaultextension=".pdf", initialname=f"compressed_{self.doc.filename}")
        if output:
            orig_size = os.path.getsize(self.doc.filepath) if self.doc.filepath else 0
            if self.doc.compress(output):
                new_size = os.path.getsize(output)
                savings = (1 - new_size / orig_size) * 100 if orig_size else 0
                messagebox.showinfo("Compressed", f"Original: {orig_size // 1024} KB\nCompressed: {new_size // 1024} KB\nSaved: {savings:.1f}%")
    
    def _ocr_doc(self):
        if not self.doc:
            return
        ok, msg = OCREngine.is_available()
        if not ok:
            messagebox.showerror("OCR Unavailable", msg)
            return
        
        if not messagebox.askyesno("OCR", "Make document searchable?\nThis may take a while."):
            return
        
        dialog = self._create_dialog("OCR Processing", 300, 100)
        label = tk.Label(dialog, text="Processing...", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY)
        label.pack(expand=True)
        
        def run():
            ok, count = OCREngine.make_searchable(self.doc, lambda m: self.after(0, lambda: label.configure(text=m)))
            self.after(0, lambda: self._ocr_done(ok, count, dialog))
        
        threading.Thread(target=run, daemon=True).start()
    
    def _ocr_done(self, ok, count, dialog):
        dialog.destroy()
        if ok:
            self._render_page()
            messagebox.showinfo("OCR Complete", f"Processed {count} pages.\nDocument is now searchable.")
        else:
            messagebox.showerror("OCR Failed", "OCR processing failed")
    
    # =========================================================================
    # EXPORTS
    # =========================================================================
    
    def _export_word(self):
        if not self.doc:
            return
        if not HAS_DOCX:
            messagebox.showerror("Unavailable", "python-docx not installed.\nRestart the app to auto-install.")
            return
        output = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if output:
            if self.doc.export_to_word(output):
                self._status("Exported to Word")
                messagebox.showinfo("Done", "Exported to Word format")
            else:
                messagebox.showerror("Error", "Export failed")
    
    def _export_images(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Export to Images", 320, 220)
        
        tk.Label(dialog, text="DPI:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_LG, Theme.PAD_XS))
        dpi_var = tk.StringVar(value="150")
        tk.Spinbox(dialog, from_=72, to=600, textvariable=dpi_var, width=10, bg=Theme.BG_INPUT, fg=Theme.FG_PRIMARY).pack()
        
        tk.Label(dialog, text="Format:", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY).pack(pady=(Theme.PAD_MD, Theme.PAD_XS))
        fmt_var = tk.StringVar(value="png")
        ttk.Combobox(dialog, textvariable=fmt_var, values=["png", "jpg"], width=10).pack()
        
        def export():
            output_dir = filedialog.askdirectory(title="Select output folder")
            if output_dir:
                files = self.doc.export_to_images(output_dir, int(dpi_var.get()), fmt_var.get())
                dialog.destroy()
                messagebox.showinfo("Done", f"Exported {len(files)} images")
        
        ModernButton(dialog, text="Export", command=export, style="primary", width=100).pack(pady=Theme.PAD_LG)
    
    def _export_text(self):
        if not self.doc:
            return
        output = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt")])
        if output:
            if self.doc.export_text(output):
                self._status("Text exported")
                messagebox.showinfo("Done", "Text extracted")
            else:
                messagebox.showerror("Error", "Export failed")
    
    # =========================================================================
    # MISC
    # =========================================================================
    
    def _show_properties(self):
        if not self.doc:
            return
        dialog = self._create_dialog("Document Properties", 420, 360)
        
        meta = self.doc.get_metadata()
        fields = [("Title", "title"), ("Author", "author"), ("Subject", "subject"), ("Keywords", "keywords")]
        entries = {}
        
        for label, key in fields:
            frame = tk.Frame(dialog, bg=Theme.BG_SECONDARY)
            frame.pack(fill=tk.X, padx=Theme.PAD_LG, pady=Theme.PAD_SM)
            tk.Label(frame, text=label + ":", bg=Theme.BG_SECONDARY, fg=Theme.FG_PRIMARY, width=12, anchor="w").pack(side=tk.LEFT)
            entry = ModernEntry(frame, width=32)
            entry.pack(side=tk.LEFT, ipady=3)
            entry.insert(0, meta.get(key, '') or '')
            entries[key] = entry
        
        tk.Label(dialog, text=f"\nFile: {self.doc.filename}", bg=Theme.BG_SECONDARY, fg=Theme.FG_MUTED).pack()
        tk.Label(dialog, text=f"Pages: {self.doc.page_count}", bg=Theme.BG_SECONDARY, fg=Theme.FG_MUTED).pack()
        if self.doc.filepath:
            tk.Label(dialog, text=f"Size: {os.path.getsize(self.doc.filepath) // 1024} KB", bg=Theme.BG_SECONDARY, fg=Theme.FG_MUTED).pack()
        
        def save():
            self.doc.set_metadata({k: e.get() for k, e in entries.items()})
            dialog.destroy()
            self._status("Properties updated")
        
        ModernButton(dialog, text="Save", command=save, style="primary", width=100).pack(pady=Theme.PAD_LG)
    
    def _show_shortcuts(self):
        shortcuts = """
Keyboard Shortcuts

FILE
  Ctrl+N    New document
  Ctrl+O    Open file
  Ctrl+S    Save
  Ctrl+W    Close tab

NAVIGATION
  Home      First page
  End       Last page
  PgUp      Previous page
  PgDn      Next page

VIEW
  Ctrl++    Zoom in
  Ctrl+-    Zoom out
  Ctrl+0    Fit page

TOOLS
  Escape    Select tool
  Delete    Delete page
  Ctrl+F    Find text
"""
        messagebox.showinfo("Keyboard Shortcuts", shortcuts)
    
    def _show_about(self):
        messagebox.showinfo("About",
            "PDF Editor Pro v4.0\n\n"
            "Professional PDF Editing Suite\n\n"
            "Features:\n"
            "• Multi-document tabs\n"
            "• Search & navigation\n"
            "• Annotations & comments\n"
            "• Stamps library\n"
            "• Watermarks & headers\n"
            "• Bates numbering\n"
            "• OCR text recognition\n"
            "• Export to Word/images\n"
            "• Merge, split, compress\n"
            "• Password protection\n\n"
            "© 2025")
    
    # =========================================================================
    # CLEANUP
    # =========================================================================
    
    def _on_close(self):
        for doc in self.documents.values():
            if doc.is_modified:
                r = messagebox.askyesnocancel("Save Changes?", f"Save changes to {doc.filename}?")
                if r is None:
                    return
                if r:
                    if doc.filepath:
                        doc.save()
                    else:
                        path = filedialog.asksaveasfilename(defaultextension=".pdf")
                        if path:
                            doc.save(path)
        
        self.config_data["window_geometry"] = self.geometry()
        Config.save(self.config_data)
        
        for doc in self.documents.values():
            doc.close()
        
        self.destroy()

# ============================================================================
# MAIN
# ============================================================================

def main():
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    # Configure ttk styles
    app = PDFEditorPro()
    
    style = ttk.Style()
    style.theme_use('clam')
    style.configure("TScrollbar", background=Theme.BG_ACTIVE, troughcolor=Theme.BG_PRIMARY,
                   bordercolor=Theme.BG_PRIMARY, arrowcolor=Theme.FG_PRIMARY)
    style.configure("TCombobox", fieldbackground=Theme.BG_INPUT, background=Theme.BG_INPUT,
                   foreground=Theme.FG_PRIMARY, selectbackground=Theme.ACCENT)
    style.configure("Treeview", background=Theme.BG_INPUT, foreground=Theme.FG_PRIMARY,
                   fieldbackground=Theme.BG_INPUT)
    style.map("Treeview", background=[("selected", Theme.ACCENT)])
    
    app.mainloop()

if __name__ == "__main__":
    main()
